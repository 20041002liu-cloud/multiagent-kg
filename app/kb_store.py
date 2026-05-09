from __future__ import annotations

import json
import shutil
import subprocess
import threading
from concurrent.futures import ThreadPoolExecutor, as_completed
from datetime import datetime, timezone
from pathlib import Path
from typing import Any
from uuid import uuid4

from docx import Document as DocxDocument
from pypdf import PdfReader

from .data_utils import split_text


def _utc_now_iso() -> str:
    return datetime.now(timezone.utc).isoformat()


class KnowledgeBaseStore:
    def __init__(
        self,
        base_dir: Path,
        *,
        pdf_ocr_enabled: bool = True,
        pdf_ocr_language: str = "chi_sim+eng",
        pdf_ocr_dpi: int = 180,
        pdf_ocr_max_pages: int = 0,
        pdf_ocr_timeout_seconds: int = 120,
        pdf_ocr_concurrency: int = 1,
        pdf_ocr_min_chars_per_page: int = 80,
        tesseract_cmd: str | None = None,
    ) -> None:
        self._base_dir = base_dir
        self._base_dir.mkdir(parents=True, exist_ok=True)
        self._meta_file = self._base_dir / "kb_store.json"
        self._docs_dir = self._base_dir / "documents"
        self._docs_dir.mkdir(parents=True, exist_ok=True)
        self._lock = threading.Lock()
        self._pdf_ocr_enabled = pdf_ocr_enabled
        self._pdf_ocr_language = pdf_ocr_language
        self._pdf_ocr_dpi = pdf_ocr_dpi
        self._pdf_ocr_max_pages = pdf_ocr_max_pages
        self._pdf_ocr_timeout_seconds = pdf_ocr_timeout_seconds
        self._pdf_ocr_concurrency = max(1, pdf_ocr_concurrency)
        self._pdf_ocr_min_chars_per_page = max(0, pdf_ocr_min_chars_per_page)
        self._tesseract_cmd = tesseract_cmd
        self._ensure_meta()

    def _ensure_meta(self) -> None:
        if self._meta_file.exists():
            return
        seed = {
            "knowledge_bases": [
                {
                    "id": "default",
                    "name": "Default KB",
                    "description": "Default knowledge base",
                    "created_at": _utc_now_iso(),
                    "updated_at": _utc_now_iso(),
                }
            ],
            "documents": [],
        }
        self._meta_file.write_text(json.dumps(seed, ensure_ascii=False, indent=2), encoding="utf-8")

    def _load(self) -> dict[str, Any]:
        return json.loads(self._meta_file.read_text(encoding="utf-8"))

    def _save(self, data: dict[str, Any]) -> None:
        self._meta_file.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")

    def list_kbs(self) -> list[dict[str, Any]]:
        with self._lock:
            data = self._load()
            return list(data["knowledge_bases"])

    def get_kb(self, kb_id: str) -> dict[str, Any] | None:
        with self._lock:
            data = self._load()
            for item in data["knowledge_bases"]:
                if item["id"] == kb_id:
                    return dict(item)
        return None

    def create_kb(self, name: str, description: str = "") -> dict[str, Any]:
        with self._lock:
            data = self._load()
            kb_id = str(uuid4())
            row = {
                "id": kb_id,
                "name": name.strip(),
                "description": description.strip(),
                "created_at": _utc_now_iso(),
                "updated_at": _utc_now_iso(),
            }
            data["knowledge_bases"].append(row)
            self._save(data)
            return row

    def delete_kb(self, kb_id: str) -> dict[str, Any]:
        if kb_id == "default":
            raise ValueError("default_kb_cannot_be_deleted")

        with self._lock:
            data = self._load()
            kb_exists = any(x["id"] == kb_id for x in data["knowledge_bases"])
            if not kb_exists:
                raise ValueError("knowledge_base_not_found")
            documents = [x for x in data["documents"] if x["knowledge_base_id"] == kb_id]
            data["knowledge_bases"] = [x for x in data["knowledge_bases"] if x["id"] != kb_id]
            data["documents"] = [x for x in data["documents"] if x["knowledge_base_id"] != kb_id]
            self._save(data)

        kb_dir = self._docs_dir / kb_id
        self._remove_path_if_safe(kb_dir)
        return {"deleted": True, "knowledge_base_id": kb_id, "document_count": len(documents)}

    def list_documents(self, kb_id: str) -> list[dict[str, Any]]:
        with self._lock:
            data = self._load()
            rows = [x for x in data["documents"] if x["knowledge_base_id"] == kb_id]
            return sorted(rows, key=lambda x: x["created_at"], reverse=True)

    def get_document(self, kb_id: str, document_id: str) -> dict[str, Any] | None:
        with self._lock:
            data = self._load()
            for row in data["documents"]:
                if row["id"] == document_id and row["knowledge_base_id"] == kb_id:
                    return dict(row)
        return None

    def _update_document(self, kb_id: str, document_id: str, patch: dict[str, Any]) -> dict[str, Any]:
        with self._lock:
            data = self._load()
            for i, row in enumerate(data["documents"]):
                if row["id"] == document_id and row["knowledge_base_id"] == kb_id:
                    updated = dict(row)
                    updated.update(patch)
                    updated["updated_at"] = _utc_now_iso()
                    data["documents"][i] = updated
                    self._save(data)
                    return updated
        raise ValueError("document not found")

    def mark_document_parsing(self, kb_id: str, document_id: str) -> dict[str, Any]:
        return self._update_document(
            kb_id,
            document_id,
            {"status": "parsing", "error": None, "chunk_count": 0},
        )

    def mark_document_failed(self, kb_id: str, document_id: str, error: str) -> dict[str, Any]:
        return self._update_document(
            kb_id,
            document_id,
            {"status": "failed", "error": error, "chunk_count": 0},
        )

    def delete_document(self, kb_id: str, document_id: str) -> dict[str, Any]:
        with self._lock:
            data = self._load()
            target: dict[str, Any] | None = None
            remaining: list[dict[str, Any]] = []
            for row in data["documents"]:
                if row["id"] == document_id and row["knowledge_base_id"] == kb_id:
                    target = row
                    continue
                remaining.append(row)
            if target is None:
                raise ValueError("document_not_found")
            data["documents"] = remaining
            self._save(data)

        for key in ("raw_path", "text_path", "chunks_path"):
            value = target.get(key)
            if value:
                self._remove_path_if_safe(Path(value))
        return {"deleted": True, "document_id": document_id}

    def _remove_path_if_safe(self, path: Path) -> None:
        try:
            resolved = path.resolve()
            docs_root = self._docs_dir.resolve()
        except OSError:
            return
        if resolved == docs_root or docs_root not in resolved.parents:
            return
        if resolved.is_dir():
            shutil.rmtree(resolved, ignore_errors=True)
        else:
            resolved.unlink(missing_ok=True)

    def create_document(self, kb_id: str, filename: str, file_bytes: bytes) -> dict[str, Any]:
        suffix = Path(filename).suffix.lower()
        if suffix not in {".pdf", ".docx", ".txt"}:
            raise ValueError(f"unsupported_file_type: {suffix}")

        doc_id = str(uuid4())
        kb_dir = self._docs_dir / kb_id
        kb_dir.mkdir(parents=True, exist_ok=True)
        raw_path = kb_dir / f"{doc_id}{suffix}"
        text_path = kb_dir / f"{doc_id}.txt"
        chunks_path = kb_dir / f"{doc_id}.chunks.json"

        raw_path.write_bytes(file_bytes)
        row = {
            "id": doc_id,
            "knowledge_base_id": kb_id,
            "filename": filename,
            "file_type": suffix.lstrip("."),
            "status": "uploaded",
            "chunk_count": 0,
            "error": None,
            "size_bytes": len(file_bytes),
            "created_at": _utc_now_iso(),
            "updated_at": _utc_now_iso(),
            "raw_path": str(raw_path),
            "text_path": str(text_path),
            "chunks_path": str(chunks_path),
        }
        with self._lock:
            data = self._load()
            data["documents"].append(row)
            self._save(data)
        return row

    def parse_document(self, kb_id: str, document_id: str, chunk_size: int, chunk_overlap: int) -> dict[str, Any]:
        doc = self.get_document(kb_id=kb_id, document_id=document_id)
        if not doc:
            raise ValueError("document_not_found")

        self._update_document(kb_id, document_id, {"status": "parsing", "error": None})
        raw_path = Path(doc["raw_path"])
        text_path = Path(doc["text_path"])
        chunks_path = Path(doc["chunks_path"])

        try:
            if doc["file_type"] == "txt":
                text = self._parse_txt(raw_path)
            elif doc["file_type"] == "docx":
                text = self._parse_docx(raw_path)
            elif doc["file_type"] == "pdf":
                text = self._parse_pdf(raw_path)
            else:
                raise ValueError("unsupported_file_type")
        except Exception as exc:
            return self._update_document(
                kb_id,
                document_id,
                {"status": "failed", "error": f"{type(exc).__name__}: {exc}", "chunk_count": 0},
            )

        normalized = self._normalize_text(text)
        if not normalized:
            return self._update_document(
                kb_id,
                document_id,
                {"status": "failed", "error": "empty_extracted_text", "chunk_count": 0},
            )

        chunks = split_text(normalized, chunk_size=chunk_size, overlap=chunk_overlap)
        text_path.write_text(normalized, encoding="utf-8")
        chunks_path.write_text(json.dumps(chunks, ensure_ascii=False, indent=2), encoding="utf-8")

        status = "parsed_low_quality" if len(normalized) < 80 else "parsed"
        return self._update_document(
            kb_id,
            document_id,
            {"status": status, "error": None, "chunk_count": len(chunks)},
        )

    def get_document_text(self, kb_id: str, document_id: str) -> str:
        doc = self.get_document(kb_id=kb_id, document_id=document_id)
        if not doc:
            raise ValueError("document_not_found")
        path = Path(doc["text_path"])
        if not path.exists():
            return ""
        return path.read_text(encoding="utf-8")

    def get_document_chunks(self, kb_id: str, document_id: str) -> list[str]:
        doc = self.get_document(kb_id=kb_id, document_id=document_id)
        if not doc:
            raise ValueError("document_not_found")
        path = Path(doc["chunks_path"])
        if not path.exists():
            return []
        try:
            data = json.loads(path.read_text(encoding="utf-8"))
        except json.JSONDecodeError:
            return []
        return [str(x) for x in data]

    @staticmethod
    def _normalize_text(text: str) -> str:
        lines = [line.strip() for line in text.replace("\r\n", "\n").split("\n")]
        cleaned = "\n".join([x for x in lines if x])
        return cleaned.strip()

    @staticmethod
    def _parse_txt(path: Path) -> str:
        for encoding in ("utf-8", "utf-8-sig", "gbk", "gb18030"):
            try:
                return path.read_text(encoding=encoding)
            except UnicodeDecodeError:
                continue
        return path.read_text(encoding="latin-1")

    @staticmethod
    def _parse_docx(path: Path) -> str:
        doc = DocxDocument(str(path))
        parts: list[str] = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                parts.append(text)
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    parts.append(row_text)
        return "\n".join(parts)

    def _parse_pdf(self, path: Path) -> str:
        text, page_count = self._parse_pdf_text(path)
        if not self._pdf_ocr_enabled:
            return text
        if text.strip() and not self._should_use_ocr_fallback(text, page_count):
            return text
        return self._parse_pdf_ocr(path)

    @staticmethod
    def _parse_pdf_text(path: Path) -> tuple[str, int]:
        reader = PdfReader(str(path))
        out: list[str] = []
        for page in reader.pages:
            text = (page.extract_text() or "").strip()
            if text:
                out.append(text)
        return "\n".join(out), len(reader.pages)

    def _should_use_ocr_fallback(self, text: str, page_count: int) -> bool:
        if page_count < 20 or self._pdf_ocr_min_chars_per_page <= 0:
            return False
        return len(text.strip()) / max(1, page_count) < self._pdf_ocr_min_chars_per_page

    def _resolve_tesseract_cmd(self) -> str | None:
        candidates: list[str] = []
        if self._tesseract_cmd:
            candidates.append(self._tesseract_cmd)
            configured = shutil.which(self._tesseract_cmd)
            if configured:
                candidates.append(configured)
        found = shutil.which("tesseract")
        if found:
            candidates.append(found)
        candidates.extend(
            [
                r"C:\Program Files\Tesseract-OCR\tesseract.exe",
                r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
            ]
        )
        for candidate in candidates:
            path = Path(candidate)
            if path.exists():
                return str(path)
        return None

    def _parse_pdf_ocr(self, path: Path) -> str:
        tesseract = self._resolve_tesseract_cmd()
        if not tesseract:
            raise RuntimeError("ocr_engine_not_found: install Tesseract OCR and set TESSERACT_CMD if needed")

        try:
            import fitz  # PyMuPDF
        except Exception as exc:  # pragma: no cover - depends on optional local package.
            raise RuntimeError(f"ocr_renderer_missing: install PyMuPDF ({type(exc).__name__}: {exc})") from exc

        zoom = max(90, self._pdf_ocr_dpi) / 72
        matrix = fitz.Matrix(zoom, zoom)

        temp_parent = self._base_dir / "ocr_tmp"
        temp_parent.mkdir(parents=True, exist_ok=True)
        tmp_dir = temp_parent / f"kg_pdf_ocr_{uuid4().hex}"
        tmp_dir.mkdir(parents=True, exist_ok=False)
        try:
            image_paths: list[tuple[int, Path]] = []
            with fitz.open(str(path)) as doc:
                page_count = len(doc)
                max_pages = self._pdf_ocr_max_pages if self._pdf_ocr_max_pages > 0 else page_count
                for page_index in range(min(page_count, max_pages)):
                    page = doc.load_page(page_index)
                    pix = page.get_pixmap(matrix=matrix, alpha=False)
                    image_path = tmp_dir / f"page_{page_index + 1:04d}.png"
                    image_path.write_bytes(pix.tobytes("png"))
                    image_paths.append((page_index + 1, image_path))

            def run_ocr(page_number: int, image_path: Path) -> tuple[int, str]:
                cmd = [
                    tesseract,
                    str(image_path),
                    "stdout",
                    "-l",
                    self._pdf_ocr_language,
                    "--psm",
                    "6",
                ]
                result = subprocess.run(
                    cmd,
                    capture_output=True,
                    text=True,
                    encoding="utf-8",
                    errors="ignore",
                    timeout=self._pdf_ocr_timeout_seconds,
                    check=False,
                )
                if result.returncode != 0:
                    message = (result.stderr or result.stdout or "unknown error").strip().replace("\n", " ")
                    raise RuntimeError(f"ocr_failed_page_{page_number}: {message[:300]}")
                return page_number, result.stdout.strip()

            page_texts: dict[int, str] = {}
            workers = min(self._pdf_ocr_concurrency, len(image_paths))
            if workers <= 1:
                for page_number, image_path in image_paths:
                    page_number, text = run_ocr(page_number, image_path)
                    page_texts[page_number] = text
            else:
                with ThreadPoolExecutor(max_workers=workers) as executor:
                    futures = [executor.submit(run_ocr, page_number, image_path) for page_number, image_path in image_paths]
                    for future in as_completed(futures):
                        page_number, text = future.result()
                        page_texts[page_number] = text

            out = [
                f"Page {page_number}:\n{text}"
                for page_number, text in sorted(page_texts.items())
                if text
            ]
        finally:
            shutil.rmtree(tmp_dir, ignore_errors=True)

        return "\n\n".join(out)
